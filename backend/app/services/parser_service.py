from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any


logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}


def _normalize_extension(file_path: str | Path) -> str:
    return Path(file_path).suffix.lower()


def _clean_text(raw_text: str) -> str:
    # Normalize OCR noise while keeping line boundaries for field extraction.
    lines = [re.sub(r"[\x00-\x1F]+", " ", line).strip() for line in raw_text.splitlines()]
    lines = [re.sub(r"\s{2,}", " ", line) for line in lines if line]
    return "\n".join(lines)


def _lazy_import_image_stack() -> tuple[Any, Any, Any]:
    from PIL import Image, ImageOps
    import pytesseract

    return Image, ImageOps, pytesseract


def _lazy_import_docx() -> Any:
    from docx import Document

    return Document


def _lazy_import_pdf() -> Any:
    from pdf2image import convert_from_path

    return convert_from_path


def image_to_text(image: Any) -> str:
    Image, ImageOps, pytesseract = _lazy_import_image_stack()
    if not isinstance(image, Image.Image):
        image = Image.open(str(image))
    grayscale = ImageOps.grayscale(image)
    processed = ImageOps.autocontrast(grayscale)
    return pytesseract.image_to_string(processed)


def pdf_to_text(file_path: str | Path) -> str:
    convert_from_path = _lazy_import_pdf()
    full_text = []
    for page in convert_from_path(str(file_path)):
        full_text.append(image_to_text(page))
    return _clean_text("\n".join(full_text))


def docx_to_text(file_path: str | Path) -> str:
    Document = _lazy_import_docx()
    document = Document(str(file_path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_cells.append(cell.text.strip())
    return _clean_text("\n".join(paragraphs + table_cells))


def extract_text(file_path: str | Path) -> str:
    suffix = _normalize_extension(file_path)
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported document type: {suffix or 'unknown'}")

    if suffix == ".pdf":
        return pdf_to_text(file_path)
    if suffix == ".docx":
        return docx_to_text(file_path)

    return _clean_text(image_to_text(file_path))


def _parse_amount(value: str | None) -> float | None:
    if not value:
        return None
    digits = re.sub(r"[^\d.]", "", value)
    if not digits:
        return None
    try:
        return float(digits)
    except ValueError:
        return None


def _parse_int(value: str | None) -> int | None:
    if not value:
        return None
    digits = re.sub(r"[^\d]", "", value)
    if not digits:
        return None
    return int(digits)


def _split_name(full_name: str | None) -> tuple[str | None, str | None]:
    if not full_name:
        return None, None
    parts = [part for part in re.split(r"\s+", full_name.strip()) if part]
    if not parts:
        return None, None
    if len(parts) == 1:
        return parts[0].title(), None
    return parts[0].title(), " ".join(parts[1:]).title()


def _line_match(lines: list[str], patterns: list[str]) -> tuple[str | None, float]:
    for index, line in enumerate(lines):
        for pattern in patterns:
            match = re.search(pattern, line, flags=re.IGNORECASE)
            if match:
                value = match.group(1).strip()
                if not value:
                    continue
                confidence = 0.9
                if index > 0 and len(line) < 8:
                    confidence = 0.75
                return value, confidence
    return None, 0.0


def _keyword_presence_confidence(text: str, keywords: list[str], fallback: float) -> float:
    lowered = text.lower()
    return fallback if any(keyword in lowered for keyword in keywords) else 0.0


def extract_fields(text: str) -> tuple[dict[str, Any], dict[str, float], list[str]]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    lowered = text.lower()

    extracted: dict[str, Any] = {}
    confidence: dict[str, float] = {}

    def capture(field: str, patterns: list[str], parser: Any = None, keywords: list[str] | None = None) -> None:
        raw_value, score = _line_match(lines, patterns)
        if raw_value is None and keywords:
            score = _keyword_presence_confidence(lowered, keywords, 0.2)
        if raw_value is None:
            confidence[field] = round(score, 2)
            return

        value = parser(raw_value) if parser else raw_value
        if value is None:
            confidence[field] = round(max(0.2, score - 0.3), 2)
            return

        extracted[field] = value
        confidence[field] = round(score, 2)

    capture(
        "full_name",
        [r"(?:applicant|borrower|customer)\s*name\s*[:\-]?\s*([A-Za-z][A-Za-z .']{2,80})"],
    )
    capture("email", [r"([\w.+-]+@[\w-]+\.[\w.-]+)"])
    capture("phone", [r"(?:mobile|phone|contact)\s*[:\-]?\s*([+\d][\d\s()-]{7,20})"])
    capture(
        "monthly_income",
        [
            r"(?:monthly\s*(?:salary|income)|salary|income)\s*[:\-]?\s*₹?([\d,]+(?:\.\d+)?)",
        ],
        parser=_parse_amount,
        keywords=["salary", "monthly income", "income"],
    )
    capture(
        "annual_income",
        [r"(?:annual\s*income|yearly\s*income)\s*[:\-]?\s*₹?([\d,]+(?:\.\d+)?)"],
        parser=_parse_amount,
        keywords=["annual income", "yearly income"],
    )
    capture(
        "loan_amount",
        [r"(?:loan\s*amount|amount\s*requested|requested\s*loan)\s*[:\-]?\s*₹?([\d,]+(?:\.\d+)?)"],
        parser=_parse_amount,
        keywords=["loan amount", "requested"],
    )
    capture("emi", [r"(?:emi|monthly\s*emi)\s*[:\-]?\s*₹?([\d,]+(?:\.\d+)?)"], parser=_parse_amount, keywords=["emi"])
    capture("loan_tenure", [r"(?:loan\s*tenure|tenure)\s*[:\-]?\s*(\d{1,3})"], parser=_parse_int)
    capture("age", [r"(?:age|applicant\s*age)\s*[:\-]?\s*(\d{2})"], parser=_parse_int)
    capture("cibil_score", [r"(?:cibil|credit\s*score)\s*[:\-]?\s*(\d{3})"], parser=_parse_int, keywords=["cibil", "credit score"])
    capture("city", [r"(?:city|town)\s*[:\-]?\s*([A-Za-z][A-Za-z .-]{1,30})"])
    capture("region", [r"(?:region|zone)\s*[:\-]?\s*([A-Za-z][A-Za-z .-]{1,20})"])
    capture("employment_type", [r"(?:employment\s*type|occupation|job\s*type)\s*[:\-]?\s*([A-Za-z][A-Za-z -]{2,30})"])
    capture("marital_status", [r"(?:marital\s*status|status)\s*[:\-]?\s*(single|married|divorced|widowed)"])
    capture("gender", [r"(?:gender|sex)\s*[:\-]?\s*(male|female|other)"])

    if "gender" not in extracted:
        if "female" in lowered:
            extracted["gender"] = "female"
            confidence["gender"] = max(confidence.get("gender", 0.0), 0.5)
        elif "male" in lowered:
            extracted["gender"] = "male"
            confidence["gender"] = max(confidence.get("gender", 0.0), 0.5)

    low_confidence_fields = [field for field, score in confidence.items() if score < 0.6]
    return extracted, confidence, low_confidence_fields


def map_to_application_schema(extracted: dict[str, Any], confidence: dict[str, float]) -> tuple[dict[str, Any], list[str]]:
    defaults_applied: list[str] = []

    first_name, last_name = _split_name(extracted.get("full_name"))
    monthly_income = extracted.get("monthly_income")
    annual_income = extracted.get("annual_income")
    loan_amount = extracted.get("loan_amount")
    emi = extracted.get("emi")
    age = extracted.get("age")
    loan_tenure = extracted.get("loan_tenure")
    cibil_score = extracted.get("cibil_score")

    mapped: dict[str, Any] = {
        "firstName": first_name or "Customer",
        "lastName": last_name or "Applicant",
        "email": extracted.get("email", "customer@example.com"),
        "phone": extracted.get("phone", "+91 9000000000"),
        "loanPurpose": "personal",
        "interestRate": 12.0,
        "existingEmis": 0.0,
        "residentialAssetsValue": 0.0,
        "commercialAssetsValue": 0.0,
        "bankBalance": 0.0,
        "totalLoans": 0,
        "activeLoans": 0,
        "closedLoans": 0,
        "missedPayments": 0,
        "creditUtilizationRatio": 0.0,
        "dependents": 0,
        "yearsOfEmployment": 0,
        "gender": str(extracted.get("gender", "other")).lower(),
        "maritalStatus": str(extracted.get("marital_status", "single")).lower(),
    }

    if monthly_income is None:
        monthly_income = 50000.0
        defaults_applied.append("monthly_income")
    if annual_income is None:
        annual_income = monthly_income * 12
        defaults_applied.append("annual_income")
    if loan_amount is None:
        loan_amount = 500000.0
        defaults_applied.append("loan_amount")
    if emi is None:
        emi = 0.0
        defaults_applied.append("emi")
    if age is None:
        age = 35
        defaults_applied.append("age")
    if loan_tenure is None:
        loan_tenure = 36
        defaults_applied.append("loan_tenure")
    if cibil_score is None:
        cibil_score = 650
        defaults_applied.append("cibil_score")

    mapped.update(
        {
            "monthlyIncome": float(monthly_income),
            "annualIncome": float(annual_income),
            "loanAmount": float(loan_amount),
            "emi": float(emi),
            "loanTenure": int(loan_tenure),
            "age": int(max(18, min(age, 70))),
            "cibilScore": int(cibil_score),
            "employmentType": str(extracted.get("employment_type", "salaried")).lower(),
            "region": str(extracted.get("region", "west")).lower(),
            "city": str(extracted.get("city", "Unknown")).strip(),
        }
    )

    # Defaults are intentionally surfaced for downstream confidence handling.
    for field in defaults_applied:
        confidence[field] = min(confidence.get(field, 0.0), 0.35)

    return mapped, defaults_applied


def parse_document(file_path: str | Path, file_name: str | None = None) -> dict[str, Any]:
    try:
        raw_text = extract_text(file_path)
        extracted, confidence, low_confidence_fields = extract_fields(raw_text)
        mapped, defaults_applied = map_to_application_schema(extracted, confidence)
        suffix = _normalize_extension(file_path)

        logger.info(
            "Document parsed | file=%s type=%s extracted_fields=%s low_confidence=%s defaults=%s",
            file_name or Path(file_path).name,
            suffix.lstrip("."),
            len(extracted),
            low_confidence_fields,
            defaults_applied,
        )

        return {
            "fileName": file_name or Path(file_path).name,
            "documentType": suffix.lstrip("."),
            "rawText": raw_text,
            "extractedData": extracted,
            "mappedData": mapped,
            "confidence": confidence,
            "lowConfidenceFields": low_confidence_fields,
            "defaultsApplied": defaults_applied,
        }
    except Exception as exc:
        logger.exception("Document parsing failed")
        raise RuntimeError(f"Could not extract required fields: {exc}") from exc
